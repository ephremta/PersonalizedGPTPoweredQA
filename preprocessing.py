from typing import List
from io import BytesIO
import re
import docx
import fitz
from functools import lru_cache
from langchain.docstore.document import Document
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores.faiss import FAISS
import logging
import os
import sys

# current_dir = os.getcwd()
# parent_dir_two_levels_back = os.path.dirname(os.path.realpath(current_dir))
# print(parent_dir_two_levels_back)
# if parent_dir_two_levels_back not in sys.path:
#     sys.path.append(parent_dir_two_levels_back)
log_format = '%(asctime)s:%(message)s'
logging.basicConfig(level=logging.INFO, format=log_format)
logger = logging.getLogger(__name__)


def parse_text_from_pdf(file_content: bytes) -> List[str]:
    pdf = fitz.open(stream=BytesIO(file_content), filetype="pdf")
    output = []
    for page_number in range(pdf.page_count):
        page = pdf.load_page(page_number)
        text = page.get_text("text")
        output.append(text)
    return output


def parse_text_from_docx(file_content: bytes) -> List[str]:
    doc = docx.Document(BytesIO(file_content))
    output = []
    for paragraph in doc.paragraphs:
        output.append(paragraph.text)
    return output


@lru_cache(maxsize=None)
def parse_document_contents(file_content: bytes, file_extension: str) -> List[str]:
    # Compile the regular expression patterns for faster matching
    merge_hyphenated_pattern = re.compile(r"(\w+)-\n(\w+)")
    fix_newlines_pattern = re.compile(r"(?<!\n\s)\n(?!\s\n)")
    remove_multiple_newlines_pattern = re.compile(r"\n\s*\n")

    def clean_text(text):
        # Merge hyphenated words
        text = merge_hyphenated_pattern.sub(r"\1\2", text)
        # Fix newlines in the middle of sentences
        text = fix_newlines_pattern.sub(" ", text.strip())
        # Remove multiple newlines
        text = remove_multiple_newlines_pattern.sub("\n\n", text)
        return text
    output = []
    if file_extension.lower() == '.pdf':
        output = parse_text_from_pdf(file_content)
    elif file_extension.lower() == '.docx':
        output = parse_text_from_docx(file_content)
    else:
        raise ValueError(
            "Unsupported file type. Only PDF and DOCX files are supported.")

    # Clean the text for all file types
    output = [clean_text(text) for text in output]
    return output


def get_last_file_extension(filename: str) -> str:
    _, file_extension = os.path.splitext(filename)
    return file_extension.lower()


def text_to_docs(texts: List[str]) -> List[Document]:
    """Converts a list of strings to a list of Documents with metadata."""
    # If the input texts are already Document objects, return them as is
    if all(isinstance(text, Document) for text in texts):
        return texts
    # Otherwise, create Document objects from the texts
    page_docs = [Document(page_content=page) for page in texts]

    # Add page numbers as metadata
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    # Split pages into chunks
    doc_chunks = []
    for doc in page_docs:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=0,
        )
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, metadata={
                    "page": doc.metadata["page"], "chunk": i}
            )
            # Add sources as metadata
            doc.metadata["source"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
            doc_chunks.append(doc)
    print(doc_chunks)

    return doc_chunks


def test_embed(api_key_val: str, pages: List[str]):
    embeddings = OpenAIEmbeddings(openai_api_key=api_key_val)
    index = FAISS.from_documents(pages, embeddings)
    print(type(index))

    return index
